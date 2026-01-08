"""
知识库工具模块
提供文档上传、向量搜索等功能
"""

import logging
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import chromadb
from chromadb.config import Settings
from docx import Document
from config.settings import VectorDBConfig, UPLOAD_DIR, CHROMA_DIR

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class KnowledgeTools:
    """知识库工具类"""
    
    def __init__(self):
        """初始化知识库和向量数据库"""
        self.upload_dir = UPLOAD_DIR
        self.chroma_dir = CHROMA_DIR
        
        # 初始化Chroma客户端
        self.client = chromadb.PersistentClient(
            path=str(self.chroma_dir),
            settings=Settings(anonymized_telemetry=False)
        )
        
        # 获取或创建集合
        self.collection = self.client.get_or_create_collection(
            name=VectorDBConfig.COLLECTION_NAME,
            metadata={"description": "文档知识库"}
        )
        
        logger.info("知识库工具初始化成功")
    
    def upload_document(self, file_path: str, filename: str, build_graph: bool = True) -> Dict[str, Any]:
        """
        上传Word文档到知识库，并可选地构建知识图谱
        
        Args:
            file_path: 文件路径
            filename: 文件名
            build_graph: 是否构建知识图谱，默认True
            
        Returns:
            上传结果字典
            
        Raises:
            Exception: 上传失败时抛出异常
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 检查文件格式
            if not filename.endswith(('.docx', '.doc')):
                raise ValueError("只支持Word文档格式 (.docx, .doc)")
            
            # 读取Word文档内容
            doc = Document(file_path)
            full_text = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        full_text.append(row_text)
            
            content = "\n".join(full_text)
            
            if not content.strip():
                raise ValueError("文档内容为空")
            
            # 分块处理文本（每500字符一块，重叠100字符）
            chunk_size = 500
            overlap_size = 100  # 重叠部分大小
            chunks = []
            for i in range(0, len(content), chunk_size - overlap_size):
                chunk = content[i:i + chunk_size]
                if chunk.strip():  # 只添加非空块
                    chunks.append(chunk)
                # 如果已经到达文本末尾，停止分块
                if i + chunk_size >= len(content):
                    break
            
            # 添加到向量数据库（使用upsert以支持重复上传）
            doc_id_prefix = filename.replace('.', '_')
            ids = [f"{doc_id_prefix}_chunk_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "filename": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                for i in range(len(chunks))
            ]
            
            # 使用 upsert 替代 add，这样可以更新已存在的文档
            self.collection.upsert(
                documents=chunks,
                ids=ids,
                metadatas=metadatas
            )
            
            logger.info(f"文档 {filename} 上传成功，共 {len(chunks)} 个文本块")
            
            result = {
                "success": True,
                "filename": filename,
                "chunks": len(chunks),
                "total_chars": len(content),
                "message": f"文档上传成功，已分割为 {len(chunks)} 个文本块"
            }
            
            # 构建知识图谱
            if build_graph:
                try:
                    from mcp_server.tools.knowledge_graph_tools import knowledge_graph_tools
                    graph_result = knowledge_graph_tools.build_graph_from_document(content, filename)
                    result["knowledge_graph"] = graph_result
                    logger.info(f"文档 {filename} 的知识图谱构建完成")
                except Exception as e:
                    logger.warning(f"知识图谱构建失败，但文档已成功上传: {str(e)}")
                    result["knowledge_graph_error"] = str(e)
            
            return result
            
        except Exception as e:
            logger.error(f"上传文档失败: {str(e)}")
            return {
                "success": False,
                "filename": filename,
                "error": str(e),
                "message": f"文档上传失败: {str(e)}"
            }
    
    def search_documents(
        self, 
        query: str, 
        top_k: int = None
    ) -> List[Dict[str, Any]]:
        """
        在知识库中搜索相关文档
        
        Args:
            query: 搜索查询
            top_k: 返回结果数量，默认使用配置值
            
        Returns:
            搜索结果列表
            
        Raises:
            Exception: 搜索失败时抛出异常
        """
        try:
            if top_k is None:
                top_k = VectorDBConfig.TOP_K
            
            # 执行向量搜索
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k
            )
            
            # 格式化结果
            formatted_results = []
            if results and results['documents'] and results['documents'][0]:
                for i, doc in enumerate(results['documents'][0]):
                    metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                    distance = results['distances'][0][i] if results['distances'] else None
                    
                    formatted_results.append({
                        "content": doc,
                        "filename": metadata.get("filename", "unknown"),
                        "chunk_index": metadata.get("chunk_index", 0),
                        "relevance_score": 1 - distance if distance else None
                    })
            
            logger.info(f"搜索 '{query}' 找到 {len(formatted_results)} 条相关结果")
            return formatted_results
            
        except Exception as e:
            logger.error(f"搜索文档失败: {str(e)}")
            raise
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """
        列出知识库中的所有文档
        
        Returns:
            文档列表
            
        Raises:
            Exception: 获取文档列表失败时抛出异常
        """
        try:
            # 获取所有文档
            all_data = self.collection.get()
            
            # 提取唯一的文件名
            filenames = set()
            if all_data and all_data['metadatas']:
                for metadata in all_data['metadatas']:
                    if 'filename' in metadata:
                        filenames.add(metadata['filename'])
            
            # 统计每个文档的信息
            documents = []
            for filename in filenames:
                # 计算该文档的块数
                chunk_count = sum(
                    1 for m in all_data['metadatas'] 
                    if m.get('filename') == filename
                )
                
                documents.append({
                    "filename": filename,
                    "chunks": chunk_count
                })
            
            logger.info(f"知识库中共有 {len(documents)} 个文档")
            return documents
            
        except Exception as e:
            logger.error(f"获取文档列表失败: {str(e)}")
            raise
    
    def delete_document(self, filename: str) -> Dict[str, Any]:
        """
        从知识库中删除文档
        
        Args:
            filename: 要删除的文件名
            
        Returns:
            删除结果字典
            
        Raises:
            Exception: 删除失败时抛出异常
        """
        try:
            # 获取所有文档
            all_data = self.collection.get()
            
            # 找到要删除的文档ID
            ids_to_delete = []
            if all_data and all_data['metadatas']:
                for i, metadata in enumerate(all_data['metadatas']):
                    if metadata.get('filename') == filename:
                        ids_to_delete.append(all_data['ids'][i])
            
            if not ids_to_delete:
                return {
                    "success": False,
                    "filename": filename,
                    "message": f"未找到文档: {filename}"
                }
            
            # 删除文档
            self.collection.delete(ids=ids_to_delete)
            
            logger.info(f"文档 {filename} 已删除，共删除 {len(ids_to_delete)} 个文本块")
            
            return {
                "success": True,
                "filename": filename,
                "deleted_chunks": len(ids_to_delete),
                "message": f"文档删除成功，共删除 {len(ids_to_delete)} 个文本块"
            }
            
        except Exception as e:
            logger.error(f"删除文档失败: {str(e)}")
            return {
                "success": False,
                "filename": filename,
                "error": str(e),
                "message": f"文档删除失败: {str(e)}"
            }
    
    def get_collection_info(self) -> Dict[str, Any]:
        """
        获取知识库集合信息
        
        Returns:
            集合信息字典
        """
        try:
            count = self.collection.count()
            return {
                "name": VectorDBConfig.COLLECTION_NAME,
                "total_chunks": count,
                "persist_directory": str(self.chroma_dir)
            }
        except Exception as e:
            logger.error(f"获取集合信息失败: {str(e)}")
            raise


# 创建全局实例
knowledge_tools = KnowledgeTools()
