"""
Web应用主程序
使用FastAPI + WebSocket实现实时聊天功能
"""

import logging
import json
import asyncio
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

from openai import OpenAI
from config.settings import LLMConfig, WebConfig, UPLOAD_DIR
from mcp_server.tools.knowledge_tools import knowledge_tools
from web_app.mcp_client import get_mcp_client
# 注意：工具导入已移除，现在通过 MCP Server 自动调用

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 创建FastAPI应用
app = FastAPI(
    title="智能工具调度系统",
    description="基于FastMCP 2.0的智能工具调度系统",
    version="1.0.0"
)

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=WebConfig.ALLOW_ORIGINS, # 允许哪些来源（域名）访问你的API。例如["*"]表示所有域名都能访问
    allow_credentials=WebConfig.ALLOW_CREDENTIALS, # 是否允许携带cookie等凭证。True表示允许。
    allow_methods=WebConfig.ALLOW_METHODS,# 允许哪些HTTP方法（如GET、POST、PUT等）跨域请求。["*"]表示全部允许。
    allow_headers=WebConfig.ALLOW_HEADERS, # 允许哪些自定义请求头跨域。["*"]表示全部允许。
)

# 挂载静态文件
static_dir = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

# 初始化OpenAI客户端
client = OpenAI(
    api_key=LLMConfig.API_KEY,
    base_url=LLMConfig.API_BASE
)

# 会话管理器
class SessionManager:
    """会话管理器 - 管理每个WebSocket的对话历史"""
    
    def __init__(self, max_history: int = 20):
        """
        初始化会话管理器
        
        Args:
            max_history: 最大保留的对话轮数（不包括system消息），默认20轮
        """
        self.sessions: Dict[int, List[Dict[str, Any]]] = {}
        self.max_history = max_history
    
    def create_session(self, session_id: int):
        """创建新会话"""
        self.sessions[session_id] = [
            {
            "role": "system",
            "content": """你是一个智能助手，可以使用各种工具来帮助用户完成任务。

            可用工具类别：
            1. 知识图谱工具：
               - list_all_entities: 列出所有实体（用户问"有哪些实体"时使用）
               - search_entities_in_graph: 搜索特定实体
               - query_entity: 查询实体详细信息
               - find_entity_path: 查找实体间的关系路径
               - get_graph_statistics: 获取图谱统计信息
               
            2. 知识库工具：搜索文档内容、列出文档
            3. 数据库工具：查询员工信息
            4. 计算工具：数学计算
            5. 时间工具：获取时间、日期计算
            6. API工具：天气查询等

            【重要】工具选择指南：
            - 用户问"有哪些实体"、"列出所有实体" → 使用 list_all_entities
            - 用户问"XX实体的信息" → 使用 query_entity
            - 用户问"搜索包含XX的实体" → 使用 search_entities_in_graph
            - 用户问"文档内容" → 使用 search_documents

            【重要】工具结果处理规范：
            - 当工具返回数据后，你必须对结果进行整理和解读
            - 对于列表数据，要总结数量、列出关键信息
            - 对于文档列表，要说明文档名称和文本块数量
            - 对于员工数据，要展示姓名、职位、薪资等关键信息
            - 对于实体列表，按类型分类展示，并说明总数
            - 使用友好的语言，让用户容易理解
            - 如果数据量大，只展示前几条并说明总数

            示例：
            - 用户问"有哪些实体？"
              → 调用 list_all_entities()
              → 整理结果："知识图谱中共有12个实体，包括：产品/服务类4个（AI视频分析系统...）、地点类2个（煤矿...）"

            请根据用户的问题，选择合适的工具来完成任务。如果需要多个步骤，可以连续调用多个工具。
            记住之前的对话内容，提供连贯的交互体验。"""
            }
        ]
        logger.info(f"创建新会话: {session_id}")
    
    def get_session(self, session_id: int) -> List[Dict[str, Any]]:
        """获取会话历史"""
        if session_id not in self.sessions:
            self.create_session(session_id)
        return self.sessions[session_id]
    
    def add_message(self, session_id: int, message: Dict[str, Any]):
        """
        添加消息到会话，自动维护历史消息数量
        
        策略：保留 system 消息 + 最近的 N 轮对话
        """
        if session_id not in self.sessions:
            self.create_session(session_id)
        
        self.sessions[session_id].append(message)
        
        # 限制历史消息数量（保留system消息）
        messages = self.sessions[session_id]
        
        # 分离 system 消息和对话消息
        system_messages = [m for m in messages if m.get("role") == "system"]
        conversation_messages = [m for m in messages if m.get("role") != "system"]
        
        # 如果对话消息超过限制，只保留最近的消息
        if len(conversation_messages) > self.max_history:
            # 保留最近的 max_history 条消息
            conversation_messages = conversation_messages[-self.max_history:]
            self.sessions[session_id] = system_messages + conversation_messages
            logger.info(f"会话 {session_id} 历史消息已裁剪，保留最近 {self.max_history} 条")
    
    def clear_session(self, session_id: int):
        """清除会话"""
        if session_id in self.sessions:
            del self.sessions[session_id]
            logger.info(f"清除会话: {session_id}")
    
    def get_session_count(self) -> int:
        """获取会话数量"""
        return len(self.sessions)


# WebSocket连接管理器
class ConnectionManager:
    """WebSocket连接管理器"""
    
    def __init__(self):
        self.active_connections: List[WebSocket] = []
    
    async def connect(self, websocket: WebSocket):
        """接受新的WebSocket连接"""
        await websocket.accept()
        self.active_connections.append(websocket)
        logger.info(f"新的WebSocket连接，当前连接数: {len(self.active_connections)}")
    
    def disconnect(self, websocket: WebSocket):
        """断开WebSocket连接"""
        self.active_connections.remove(websocket)
        logger.info(f"WebSocket断开连接，当前连接数: {len(self.active_connections)}")
    
    async def send_message(self, message: dict, websocket: WebSocket):
        """发送消息到指定WebSocket"""
        await websocket.send_json(message)


manager = ConnectionManager()
session_manager = SessionManager()


async def get_tools_definition() -> List[Dict[str, Any]]:
    """
    从 MCP Server 获取工具定义（通过 HTTP 协议）
    
    优势：Web App 和 MCP Server 完全解耦，可以独立部署
    
    Returns:
        工具定义列表
    """
    try:
        mcp_client = get_mcp_client()
        tools = await mcp_client.get_tools_definition()
        logger.info(f"从 MCP Server 加载了 {len(tools)} 个工具定义")
        return tools
    except Exception as e:
        logger.error(f"从 MCP Server 加载工具失败: {e}，使用备用定义")
        # 备用：手动定义少量核心工具
        return []


async def execute_tool(tool_name: str, tool_args: Dict[str, Any]) -> Any:
    """
    执行工具调用 - 通过 MCP Client 调用 MCP Server 上的工具
    
    优势：
    1. 完全解耦 - Web App 和 MCP Server 独立部署
    2. 自动路由 - 添加新工具只需在 server.py 注册
    3. 标准协议 - 使用 MCP 标准通信协议
    
    Args:
        tool_name: 工具名称
        tool_args: 工具参数
        
    Returns:
        工具执行结果
    """
    try:
        mcp_client = get_mcp_client()
        result = await mcp_client.call_tool(tool_name, tool_args)
        return result
    except Exception as e:
        logger.error(f"工具执行失败 {tool_name}: {str(e)}")
        return {"error": str(e)}


async def process_message(user_message: str, websocket: WebSocket, session_id: int):
    """
    处理用户消息，调用大模型和工具
    
    Args:
        user_message: 用户消息
        websocket: WebSocket连接
        session_id: 会话ID
    """
    try:
        # 获取会话历史
        messages = session_manager.get_session(session_id)
        
        # 添加用户消息
        user_msg = {
            "role": "user",
            "content": user_message
        }
        session_manager.add_message(session_id, user_msg)
        
        # 获取工具定义（异步）
        tools = await get_tools_definition()
        
        # 调用大模型
        response = client.chat.completions.create(
            model=LLMConfig.MODEL_NAME,
            messages=messages,
            tools=tools,
            tool_choice="auto",
            temperature=LLMConfig.TEMPERATURE,
            max_tokens=LLMConfig.MAX_TOKENS
        )
        
        response_message = response.choices[0].message
        tool_calls = response_message.tool_calls
        
        # 如果有工具调用
        if tool_calls:
            # 发送工具调用信息
            for tool_call in tool_calls:
                tool_name = tool_call.function.name
                tool_args = json.loads(tool_call.function.arguments)
                
                await manager.send_message({
                    "type": "tool_call",
                    "tool_name": tool_name,
                    "tool_args": tool_args
                }, websocket)
                
                # 执行工具
                tool_result = await execute_tool(tool_name, tool_args)
                
                # 发送工具结果
                await manager.send_message({
                    "type": "tool_result",
                    "tool_name": tool_name,
                    "result": tool_result
                }, websocket)
                
                # 将工具结果添加到消息历史和会话
                tool_call_msg = {
                    "role": "assistant",
                    "content": None,
                    "tool_calls": [tool_call.model_dump()]
                }
                tool_result_msg = {
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "name": tool_name,
                    "content": json.dumps(tool_result, ensure_ascii=False)
                }
                session_manager.add_message(session_id, tool_call_msg)
                session_manager.add_message(session_id, tool_result_msg)
            
            # 添加额外的提示，让大模型整理结果
            format_hint = {
                "role": "system",
                "content": "现在请根据工具返回的结果，用友好、清晰的语言总结和整理信息，让用户容易理解。"
            }
            messages = session_manager.get_session(session_id)
            messages.append(format_hint)
            
            # 再次调用大模型，生成最终回复
            final_response = client.chat.completions.create(
                model=LLMConfig.MODEL_NAME,
                messages=messages,
                temperature=LLMConfig.TEMPERATURE,
                max_tokens=LLMConfig.MAX_TOKENS
            )
            
            assistant_message = final_response.choices[0].message.content
            
            # 移除临时提示，保存助手回复
            messages.pop()  # 移除格式提示
            session_manager.add_message(session_id, {
                "role": "assistant",
                "content": assistant_message
            })
        else:
            # 没有工具调用，直接返回回复
            assistant_message = response_message.content
            # 保存助手回复到会话
            session_manager.add_message(session_id, {
                "role": "assistant",
                "content": assistant_message
            })
        
        # 发送最终回复
        await manager.send_message({
            "type": "assistant_message",
            "content": assistant_message
        }, websocket)
        
    except Exception as e:
        logger.error(f"处理消息失败: {str(e)}")
        await manager.send_message({
            "type": "error",
            "content": f"处理消息失败: {str(e)}"
        }, websocket)


# ========================
# API路由
# ========================

@app.get("/", response_class=HTMLResponse)
async def get_index():
    """返回主页"""
    index_file = static_dir / "index.html"
    if index_file.exists():
        return FileResponse(index_file)
    return HTMLResponse("<h1>智能工具调度系统</h1><p>请创建 static/index.html 文件</p>")


@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket端点"""
    await manager.connect(websocket)
    
    # 使用WebSocket对象的ID作为会话ID
    session_id = id(websocket)
    session_manager.create_session(session_id)
    
    try:
        while True:
            # 接收消息
            data = await websocket.receive_json()
            message_type = data.get("type")
            
            if message_type == "user_message":
                user_message = data.get("content", "")
                # 处理用户消息
                await process_message(user_message, websocket, session_id)
            
            elif message_type == "clear_history":
                # 清除对话历史
                session_manager.clear_session(session_id)
                session_manager.create_session(session_id)
                await manager.send_message({
                    "type": "system_message",
                    "content": "对话历史已清除"
                }, websocket)
            
    except WebSocketDisconnect:
        manager.disconnect(websocket)
        session_manager.clear_session(session_id)
    except Exception as e:
        logger.error(f"WebSocket错误: {str(e)}")
        manager.disconnect(websocket)
        session_manager.clear_session(session_id)


@app.post("/upload")
async def upload_document(file: UploadFile = File(...), build_graph: bool = True):
    """
    上传文档到知识库
    
    Args:
        file: 上传的文件
        build_graph: 是否构建知识图谱，默认True
        
    Returns:
        上传结果
    """
    try:
        # 检查文件格式
        if not file.filename.endswith(('.docx', '.doc')):
            raise HTTPException(status_code=400, detail="只支持Word文档格式")
        
        logger.info(f"开始上传文档: {file.filename}，构建知识图谱: {build_graph}")
        
        # 保存文件
        file_path = UPLOAD_DIR / file.filename
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # 读取文档内容
        from docx import Document
        doc = Document(file_path)
        full_text = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    full_text.append(row_text)
        
        content_text = "\n".join(full_text)
        
        # 1. 上传到知识库（向量数据库）
        logger.info(f"正在上传到向量数据库...")
        result = knowledge_tools.upload_document(str(file_path), file.filename, build_graph=False)
        
        # 2. 通过 MCP 调用构建知识图谱（在 MCP Server 进程中）
        if build_graph:
            try:
                logger.info(f"正在构建知识图谱...")
                mcp_client = get_mcp_client()
                graph_result = await mcp_client.call_tool(
                    "build_knowledge_graph",
                    {
                        "content": content_text,
                        "filename": file.filename
                    }
                )
                result["knowledge_graph"] = graph_result
                logger.info(f"文档 {file.filename} 的知识图谱构建完成：{graph_result}")
            except Exception as e:
                logger.warning(f"知识图谱构建失败: {str(e)}")
                result["knowledge_graph_error"] = str(e)
        else:
            logger.info(f"跳过知识图谱构建")
        
        return result
        
    except Exception as e:
        logger.error(f"文件上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents")
async def get_documents():
    """
    获取知识库中的文档列表
    
    Returns:
        文档列表
    """
    try:
        result = knowledge_tools.list_documents()
        return {"success": True, "documents": result}
    except Exception as e:
        logger.error(f"获取文档列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health_check():
    """健康检查"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }


@app.get("/knowledge-graph")
async def knowledge_graph_page():
    """返回知识图谱可视化页面"""
    graph_file = static_dir / "knowledge-graph.html"
    if graph_file.exists():
        return FileResponse(graph_file)
    return HTMLResponse("<h1>知识图谱页面未找到</h1>")


@app.get("/knowledge-graph-simple")
async def knowledge_graph_simple_page():
    """返回知识图谱可视化页面（简化版，无需外部依赖）"""
    graph_file = static_dir / "knowledge-graph-simple.html"
    if graph_file.exists():
        return FileResponse(graph_file)
    return HTMLResponse("<h1>知识图谱页面未找到</h1>")


@app.get("/api/knowledge-graph/export")
async def export_knowledge_graph():
    """导出知识图谱数据"""
    try:
        mcp_client = get_mcp_client()
        result = await mcp_client.call_tool("export_knowledge_graph", {})
        return result
    except Exception as e:
        logger.error(f"导出知识图谱失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/knowledge-graph/statistics")
async def get_graph_statistics():
    """获取知识图谱统计信息"""
    try:
        mcp_client = get_mcp_client()
        result = await mcp_client.call_tool("get_graph_statistics", {})
        return result
    except Exception as e:
        logger.error(f"获取统计信息失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/knowledge-graph/entity/{entity_name}")
async def query_graph_entity(entity_name: str):
    """查询实体信息"""
    try:
        mcp_client = get_mcp_client()
        result = await mcp_client.call_tool("query_entity", {"entity_name": entity_name})
        return result
    except Exception as e:
        logger.error(f"查询实体失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def run_web_app():
    """运行Web应用"""
    logger.info(f"启动Web应用: http://{WebConfig.HOST}:{WebConfig.PORT}")
    uvicorn.run(
        "web_app.main:app",
        host=WebConfig.HOST,
        port=WebConfig.PORT,
        reload=WebConfig.RELOAD
    )


if __name__ == "__main__":
    run_web_app()
